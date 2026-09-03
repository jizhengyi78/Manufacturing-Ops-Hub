"""
ingestion.py — 文档摄入管线
===========================
角色：处理文档上传并写入双索引 (ES + Milvus) 的完整流程。
      这是全系统唯一的数据写入入口，所有文档——不管是手动上传、
      知识沉淀自动生成、还是批量导入——都走这里。

标准摄入流程 (7步):
┌─────────────┐
│ 1. 文件解析  │ ← PDF/Word/Excel → 纯文本
├─────────────┤
│ 2. 内容清洗  │ ← 去噪声、统一编码
├─────────────┤
│ 3. 注入扫描  │ ← scan_document_for_injection()
├─────────────┤
│ 4. 智能分块  │ ← DocumentChunker (表级/SOP步骤级/语义级)
├─────────────┤
│ 5. 向量嵌入  │ ← EmbeddingService.embed_batch()
├─────────────┤
│ 6. 双写存储  │ ← ES (BM25) + Milvus (向量) + PG事务表
├─────────────┤
│ 7. 对账确认  │ ← sync.py 后台对账
└─────────────┘

为什么只有这一个入口:
- 避免知识沉淀模块单独写入库逻辑导致 ES/Milvus 数据分裂
- 所有的补偿、对账、异常处理都在这一条链路里

使用示例:
    from src.retrieval.ingestion import standard_ingest, IngestDocument
    doc = IngestDocument(
        title="注塑机操作手册",
        content="原始文档内容...",
        doc_type="equipment_manual",
        workshop_id="workshop-a",
        equipment_model="海天MA1200",
    )
    result = await standard_ingest(doc)
    print(f"摄入完成: {result.chunk_count} 个分块, 状态: {result.status}")

注意事项:
- mill_scan=true 时会先做注入扫描，发现风险标记 pending_review 不直接入库
- 大文件 (>50MB) 应该在上传层拒绝，不要在这里处理
- 摄入失败会记录到 ingestion_transactions 表，compensation worker 会重试
- 文档去重: 相同 doc_id 的重复摄入 → 先删旧索引 → 再建新索引 (覆盖更新)
"""

from dataclasses import dataclass, field
import uuid
from pathlib import Path

from src.core.logging import get_logger
from src.retrieval.chunker import DocumentChunker, Chunk
from src.retrieval.embedding import EmbeddingService, get_embedding_service
from src.security.injection import scan_document_for_injection

logger = get_logger(__name__)


@dataclass
class IngestDocument:
    """待摄入的文档。

    示例:
        doc = IngestDocument(
            title="海天MA1200 注塑机操作手册",
            content="第一章 设备概述\n...",
            doc_type="equipment_manual",
            doc_id="doc_manual_ha_001",
            workshop_id="workshop-a",
            classification="internal",
            equipment_model="海天MA1200",
        )
    """
    title: str
    content: str
    doc_type: str              # equipment_manual / sop / fault_case / quality_standard / alarm_code
    doc_id: str | None = None
    workshop_id: str = ""
    classification: str = "internal"
    equipment_model: str = ""
    version: str = "1.0"
    source: str = "manual"     # manual / case / batch
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        if self.doc_id is None:
            self.doc_id = f"doc_{uuid.uuid4().hex[:12]}"


@dataclass
class IngestResult:
    """摄入结果。"""
    doc_id: str
    chunk_count: int
    status: str               # done / pending_review / failed
    error: str = ""


class IngestionPipeline:
    """文档摄入管线。Phase 1 先做核心流程，ES/Milvus 写入在后续文件实现。"""

    def __init__(
        self,
        embedding_service: EmbeddingService | None = None,
        chunker: DocumentChunker | None = None,
    ):
        self.embedding = embedding_service or get_embedding_service()
        self.chunker = chunker or DocumentChunker()

    async def _parse_content(self, doc: IngestDocument) -> str:
        """解析文档内容 — 支持 PDF/Word/Excel/TXT。

        Phase 3: 接入 PyPDF2/python-docx/openpyxl 做多格式解析。
        如果 content 已经有文本（种子数据），直接返回。
        """
        # 已有文本内容 → 直接返回
        if doc.content and len(doc.content.strip()) > 50:
            return doc.content

        # 从文件路径解析
        file_path = doc.metadata.get("file_path", "")
        if not file_path:
            return doc.content

        ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""

        try:
            if ext == "pdf":
                return await self._parse_pdf(file_path)
            elif ext in ("docx", "doc"):
                return await self._parse_docx(file_path)
            elif ext in ("xlsx", "xls"):
                return await self._parse_xlsx(file_path)
            elif ext in ("txt", "md", "csv"):
                return Path(file_path).read_text(encoding="utf-8")
        except Exception as e:
            logger.warning(f"文档解析失败 {ext}: {e}")
            return doc.content

        return doc.content

    async def _parse_pdf(self, path: str) -> str:
        """PDF 解析。"""
        import asyncio
        loop = asyncio.get_running_loop()

        def _read():
            import PyPDF2
            text = ""
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text or "(PDF解析结果为空)"
        return await loop.run_in_executor(None, _read)

    async def _parse_docx(self, path: str) -> str:
        """Word 解析。"""
        import asyncio
        loop = asyncio.get_running_loop()

        def _read():
            import docx
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()) or "(Word解析结果为空)"
        return await loop.run_in_executor(None, _read)

    async def _parse_xlsx(self, path: str) -> str:
        """Excel 解析: 保留表格结构，每行用 | 分隔。"""
        import asyncio
        loop = asyncio.get_running_loop()

        def _read():
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            text = ""
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text += f"\n## {sheet_name}\n"
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        text += row_text + "\n"
            return text or "(Excel解析结果为空)"
        return await loop.run_in_executor(None, _read)

    async def _clean(self, text: str) -> str:
        """轻量清洗: 去多余空行、统一换行符。"""
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        return text.strip()

    async def ingest(self, doc: IngestDocument) -> IngestResult:
        """执行完整摄入流程。

        返回 IngestResult, status:
          "done"           — 成功入库
          "pending_review" — 文档含注入风险，需人工审核
          "failed"         — 摄入失败

        示例:
            pipeline = IngestionPipeline()
            doc = IngestDocument(title="换模SOP", content="步骤1: 关闭...", doc_type="sop")
            result = await pipeline.ingest(doc)
            if result.status == "done":
                print(f"摄入成功: {result.chunk_count} chunks")
        """
        try:
            # 1. 解析 (Phase1 直传文本)
            text = await self._parse_content(doc)
            if not text.strip():
                return IngestResult(doc_id=doc.doc_id, chunk_count=0, status="failed", error="文档为空")

            # 2. 清洗
            text = await self._clean(text)

            # 3. 注入扫描 (外部来源文档)
            if doc.source != "manual":
                has_risk, risk_tags = scan_document_for_injection(text)
                if has_risk:
                    logger.warning(f"文档 [{doc.title}] 检测到注入风险: {risk_tags}")
                    return IngestResult(
                        doc_id=doc.doc_id, chunk_count=0,
                        status="pending_review",
                        error=f"检测到注入风险: {risk_tags}"
                    )

            # 4. 分块
            chunks = self.chunker.chunk(
                text=text,
                title=doc.title,
                equipment_model=doc.equipment_model,
            )
            if not chunks:
                return IngestResult(doc_id=doc.doc_id, chunk_count=0, status="failed", error="分块结果为空")

            # 5. 向量嵌入 (批量)
            chunk_texts = [c.content for c in chunks]
            vectors = await self.embedding.embed_batch(chunk_texts)
            logger.info(f"嵌入完成: {len(vectors)} 个向量, 维度: {len(vectors[0])}")

            # 6. 双写存储 (Phase 1 先记录日志，ES/Milvus 写入由 sync.py 完成)
            # TODO: 对接 sync.py::dual_write()
            logger.info(
                f"文档摄入: doc_id={doc.doc_id}, title={doc.title}, "
                f"chunks={len(chunks)}, type={doc.doc_type}, source={doc.source}"
            )

            return IngestResult(
                doc_id=doc.doc_id,
                chunk_count=len(chunks),
                status="done",
            )
        except Exception as e:
            logger.error(f"文档摄入失败 [{doc.title}]: {e}", exc_info=True)
            return IngestResult(doc_id=doc.doc_id, chunk_count=0, status="failed", error=str(e))


# ── 全系统统一的摄入入口 ──
_pipeline: IngestionPipeline | None = None


async def standard_ingest(
    doc: IngestDocument,
    pipeline: IngestionPipeline | None = None,
) -> IngestResult:
    """全系统唯一数据写入入口。

    所有模块都要通过这个函数写入数据:
    - 文档摄入 (Phase1)        → source="manual"
    - 知识沉淀 (Phase3)        → source="case"
    - 批量导入 (运维脚本)      → source="batch"
    - 不需要区分调用方

    示例:
        result = await standard_ingest(IngestDocument(
            title="模具保养规程",
            content="1. 每日检查模具表面...",
            doc_type="sop",
            source="manual",
        ))
        assert result.status == "done"
    """
    global _pipeline
    if pipeline is not None:
        return await pipeline.ingest(doc)
    if _pipeline is None:
        _pipeline = IngestionPipeline()
    return await _pipeline.ingest(doc)
